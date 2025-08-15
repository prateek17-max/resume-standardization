from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_template():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Personal Info Section
    doc.add_paragraph("${personal_info.name}").bold = True
    contact = doc.add_paragraph()
    contact.add_run("${personal_info.email} | ${personal_info.phone}\n")
    contact.add_run("${personal_info.location}")
    if "${personal_info.linkedin}":
        contact.add_run(" | LinkedIn: ${personal_info.linkedin}")

    # Summary Section
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph("${summary}")

    # Experience Section
    doc.add_heading('WORK EXPERIENCE', level=1)
    doc.add_paragraph(
        "{% for exp in experience %}\n"
        "${exp.title}\n"
        "${exp.company} | ${exp.location} | ${exp.dates}\n"
        "{% for resp in exp.responsibilities %}\n"
        "• ${resp}\n"
        "{% endfor %}\n"
        "{% endfor %}"
    )

    # Education Section
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph(
        "{% for edu in education %}\n"
        "${edu.degree}\n"
        "${edu.institution} | ${edu.location} | ${edu.dates}\n"
        "{% if edu.gpa %}GPA: ${edu.gpa}{% endif %}\n"
        "{% endfor %}"
    )

    # Skills Section
    doc.add_heading('SKILLS', level=1)
    doc.add_paragraph("Technical Skills:")
    doc.add_paragraph(
        "{% for skill in skills.technical %}"
        "• ${skill}\n"
        "{% endfor %}"
    )
    doc.add_paragraph("Soft Skills:")
    doc.add_paragraph(
        "{% for skill in skills.soft %}"
        "• ${skill}\n"
        "{% endfor %}"
    )

    # Certifications Section
    doc.add_paragraph(
        "{% if certifications %}\n"
        "CERTIFICATIONS\n"
        "{% for cert in certifications %}"
        "• ${cert}\n"
        "{% endfor %}\n"
        "{% endif %}"
    )

    # Projects Section
    doc.add_paragraph(
        "{% if projects %}\n"
        "PROJECTS\n"
        "{% for project in projects %}"
        "• ${project}\n"
        "{% endfor %}\n"
        "{% endif %}"
    )

    # Create templates directory if it doesn't exist
    os.makedirs('templates', exist_ok=True)
    
    # Save the template
    template_path = os.path.join('templates', 'standard_resume_template.docx')
    doc.save(template_path)
    return template_path

if __name__ == "__main__":
    create_template()
